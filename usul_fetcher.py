import argparse
import io
import json
import re
import sys
import time
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import Flask components
from flask import Flask, render_template_string, request, jsonify, send_file

app = Flask(__name__)

# Terminal styling constants
CLR_HEADER = "\033[95m"
CLR_BLUE = "\033[94m"
CLR_GREEN = "\033[92m"
CLR_YELLOW = "\033[93m"
CLR_FAIL = "\033[91m"
CLR_RESET = "\033[0m"
CLR_BOLD = "\033[1m"

STATUS_OK = f"[{CLR_GREEN}✓{CLR_RESET}]"
STATUS_INFO = f"[{CLR_BLUE}i{CLR_RESET}]"
STATUS_WARN = f"[{CLR_YELLOW}!{CLR_RESET}]"
STATUS_ERROR = f"[{CLR_FAIL}✗{CLR_RESET}]"

# HTML Frontend Template - Dark Theme & Brown Color Scheme
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Usul Book Fetcher</title>
    <script src="https://cdn.jsdelivr.net/npm/@tailwindcss/browser@4"></script>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Amiri:ital,wght@0,400;0,700;1,400;1,700&family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        body { font-family: 'Inter', sans-serif; }
    </style>
</head>
<body class="bg-stone-950 min-h-screen flex flex-col justify-between text-stone-200">

    <header class="bg-stone-900 border-b border-stone-800 py-5 shadow-md">
        <div class="max-w-4xl mx-auto px-4 flex items-center justify-between">
            <div class="flex items-center space-x-3">
                <div class="bg-amber-700 text-stone-100 p-2 rounded-lg font-bold text-lg tracking-wider shadow-inner">Book Fetcher</div>
                <div>
                    <h1 class="text-xl font-bold text-amber-100 tracking-wide">Usul Book Fetcher</h1>
                    <p class="text-xs text-stone-400">JSON Parser & Document Creation Engine</p>
                </div>
            </div>
            <span class="inline-flex items-center gap-x-1.5 py-1.5 px-3 rounded-full text-xs font-medium bg-amber-950/50 text-amber-400 border border-amber-800/50">
                <span class="size-1.5 inline-block rounded-full bg-amber-500 animate-pulse"></span> Engine Active
            </span>
        </div>
    </header>

    <main class="max-w-4xl mx-auto px-4 py-10 flex-1 w-full">
        <div class="bg-stone-900 rounded-2xl shadow-xl border border-stone-800 p-6 md:p-8 space-y-8">
            
            <div class="border-l-4 border-amber-600 bg-amber-950/20 p-4 rounded-r-xl border-t border-b border-r border-stone-800">
                <h3 class="text-sm font-semibold text-amber-400">How to use</h3>
                <p class="text-xs text-stone-400 mt-1">Paste URL of the required book (e.g., <code class="bg-stone-950 text-amber-300 px-1 py-0.5 rounded border border-stone-800">https://usul.ai/t/book</code>). The system will fetch the book and compile a Word Document (`.docx`).</p>
            </div>

            <form id="downloadForm" class="space-y-6">
                <div class="grid grid-cols-1 md:grid-cols-3 gap-4">
                    <div class="md:col-span-2">
                        <label for="url" class="block text-sm font-medium text-stone-300 mb-1">Usul.ai Book URL</label>
                        <input type="url" id="url" name="url" required placeholder="https://usul.ai/t/your-book" 
                               class="w-full px-4 py-2.5 rounded-xl bg-stone-950 border border-stone-800 text-stone-100 placeholder-stone-600 focus:ring-2 focus:ring-amber-500/20 focus:border-amber-600 outline-none transition-all text-sm">
                    </div>
                    <div>
                        <label for="output" class="block text-sm font-medium text-stone-300 mb-1">Output File Name</label>
                        <input type="text" id="output" name="output" value="book.docx" placeholder="book.docx"
                               class="w-full px-4 py-2.5 rounded-xl bg-stone-950 border border-stone-800 text-stone-100 placeholder-stone-600 focus:ring-2 focus:ring-amber-500/20 focus:border-amber-600 outline-none transition-all text-sm">
                    </div>
                </div>

                <button type="submit" id="submitBtn" 
                        class="w-full bg-amber-700 hover:bg-amber-600 text-stone-100 font-medium py-3 px-4 rounded-xl shadow-lg shadow-amber-900/20 transition-all flex items-center justify-center space-x-2 cursor-pointer border border-amber-600/30">
                    <span>Fetch and Compile Document</span>
                </button>
            </form>

            <div id="statusPanel" class="hidden border border-stone-800 rounded-xl overflow-hidden bg-stone-950 text-stone-300 font-mono text-xs">
                <div class="bg-stone-900 px-4 py-2 flex items-center justify-between border-b border-stone-800">
                    <span class="text-stone-400">Execution Pipeline Console</span>
                    <span id="pipelineState" class="text-amber-500 font-semibold animate-pulse">PROCESSING...</span>
                </div>
                <div id="consoleLogs" class="p-4 h-48 overflow-y-auto space-y-1.5 selection:bg-amber-700 selection:text-stone-100">
                    </div>
            </div>

        </div>
    </main>

    <footer class="bg-stone-900 border-t border-stone-800 py-4 text-center text-xs text-stone-500">
        <p>Usul Book Fetcher</p>
    </footer>

    <script>
        document.getElementById('downloadForm').addEventListener('submit', async (e) => {
            e.preventDefault();

            const urlInput = document.getElementById('url').value;
            const outputInput = document.getElementById('output').value;
            
            const submitBtn = document.getElementById('submitBtn');
            const statusPanel = document.getElementById('statusPanel');
            const consoleLogs = document.getElementById('consoleLogs');
            const pipelineState = document.getElementById('pipelineState');

            submitBtn.disabled = true;
            submitBtn.classList.add('opacity-50', 'cursor-not-allowed');
            statusPanel.classList.remove('hidden');
            consoleLogs.innerHTML = '';
            
            function appendLog(text, colorClass = 'text-stone-300') {
                const p = document.createElement('p');
                p.className = colorClass;
                p.innerHTML = `[${new Date().toLocaleTimeString()}] ${text}`;
                consoleLogs.appendChild(p);
                consoleLogs.scrollTop = consoleLogs.scrollHeight;
            }

            appendLog("Initiating core fetch engine pipeline...", "text-amber-400");

            try {
                const response = await fetch('/api/download', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ url: urlInput, output: outputInput })
                });

                if (!response.ok) {
                    const errorData = await response.json();
                    throw new Error(errorData.error || 'Server parsing pipeline failure.');
                }

                const blob = await response.blob();
                appendLog("Data streams fully captured and parsed.", "text-emerald-400");
                appendLog("Structural formatting complete.", "text-emerald-400");
                
                pipelineState.innerText = "SUCCESS";
                pipelineState.className = "text-emerald-400 font-semibold";
                
                const downloadUrl = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = downloadUrl;
                a.download = outputInput || "book.docx";
                document.body.appendChild(a);
                a.click();
                a.remove();
                
                appendLog("Document compiled and delivered successfully.", "text-emerald-400 font-bold");

            } catch (err) {
                appendLog(`Error: ${err.message}`, "text-rose-400 font-bold");
                pipelineState.innerText = "CRASHED";
                pipelineState.className = "text-rose-400 font-semibold";
            } finally {
                submitBtn.disabled = false;
                submitBtn.classList.remove('opacity-50', 'cursor-not-allowed');
            }
        });
    </script>
</body>
</html>
"""

def print_banner():
    print(f"\n{CLR_BOLD}{CLR_HEADER}==============================================")
    print("               USUL BOOK FETCHER              ")
    print(f"=============================================={CLR_RESET}\n")


def clean_and_format_url(input_url):
    match = re.search(r"/(?:t|book)/([^/?#]+)", input_url)
    if not match:
        raise ValueError(f"Could not parse book slug from URL '{input_url}'. Expected format: https://usul.ai/t/book-slug")
    book_slug = match.group(1)
    return f"https://api.usul.ai/book/{book_slug}"


def fetch_entire_book(base_url, chunk_size=50):
    all_data = []
    start_index = 0

    print(f"{STATUS_INFO} Target API Endpoint: {CLR_BLUE}{base_url}{CLR_RESET}")
    print(f"{STATUS_INFO} Initializing data stream (Chunk Size: {chunk_size})...\n")

    while True:
        params = {"startIndex": start_index, "size": chunk_size}
        try:
            sys.stdout.write(f"\r{CLR_BOLD}{CLR_BLUE} 📥 Fetching records:{CLR_RESET} {start_index:<3}")
            sys.stdout.flush()

            response = requests.get(base_url, params=params, timeout=15)
            response.raise_for_status()
            data = response.json()

            if not data:
                print(f"\r{STATUS_OK} Reached end of data stream {CLR_GREEN}(Empty response received).{CLR_RESET}")
                break

            if isinstance(data, dict):
                pages = data.get("content", {}).get("pages")
                if not pages:
                    print(f"\r{STATUS_OK} Reached end of text {CLR_GREEN}(Inner content field is empty).{CLR_RESET}")
                    break
                all_data.append(data)

            elif isinstance(data, list):
                if len(data) == 0:
                    break
                first_item = data[0]
                if isinstance(first_item, dict) and "pages" in first_item:
                    if not first_item["pages"]:
                        print(f"\r{STATUS_OK} Reached end of book structure {CLR_GREEN}('pages' array is empty).{CLR_RESET}")
                        break
                all_data.extend(data)
                if len(data) < chunk_size:
                    print(f"\r{STATUS_OK} Finished reading data stream {CLR_GREEN}(Final partial batch received).{CLR_RESET}")
                    break

            start_index += chunk_size
            time.sleep(0.5)

        except requests.exceptions.RequestException as e:
            print(f"\n\n{STATUS_ERROR} {CLR_FAIL}Network error occurred:{CLR_RESET} {e}")
            break

    return all_data


def apply_arabic_font(run, font_name, size_pt, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.complex_script = True

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

    size_val = str(int(size_pt * 2))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), size_val)
    rPr.append(szCs)

    if bold:
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)


def build_docx_from_data(data):
    doc = Document()
    text_found = False
    span_regex = re.compile(r'<span[^>]*>(.*?)</span>')
    total_chunks = len(data)

    for idx, item in enumerate(data, start=1):
        percent = int((idx / total_chunks) * 100)
        sys.stdout.write(f"\r ⚙️  Formatting Document Layout: [{percent:3}%] Processing chunk {idx}/{total_chunks}...")
        sys.stdout.flush()

        pages = item.get("content", {}).get("pages", [])
        for page in pages:
            text = page.get("text", "")
            if not text:
                continue

            text_found = True
            lines = text.split('\n')

            for line in lines:
                cleaned_line = line.strip()
                if not cleaned_line:
                    continue

                span_match = span_regex.search(cleaned_line)
                if span_match:
                    title_text = span_match.group(1).strip()
                    if title_text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(title_text)
                        apply_arabic_font(run, font_name='Noto Kufi Arabic', size_pt=18, bold=True)
                else:
                    body_text = re.sub(r'<[^>]+>', '', cleaned_line).strip()
                    if body_text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run = p.add_run(body_text)
                        apply_arabic_font(run, font_name='Scheherazade New', size_pt=14, bold=False)

    if not text_found:
        return None
        
    return doc


# Flask Server Web Routes
@app.route('/')
def home():
    return render_template_string(HTML_TEMPLATE)


@app.route('/api/download', methods=['POST'])
def handle_download():
    payload = request.get_json() or {}
    input_url = payload.get('url')
    output_filename = payload.get('output', 'book.docx') or 'book.docx'

    if not input_url:
        return jsonify({'error': 'Missing book configuration target URL.'}), 400

    try:
        api_url = clean_and_format_url(input_url)
        complete_book_data = fetch_entire_book(api_url, chunk_size=50)
        
        if not complete_book_data:
            return jsonify({'error': 'Extraction pipeline generated no usable elements.'}), 400
            
        doc = build_docx_from_data(complete_book_data)
        
        if not doc:
            return jsonify({'error': 'Valid text tags missing from data structure blocks.'}), 400

        # Save document straight to RAM instead of local storage
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        )

    except ValueError as val_err:
        return jsonify({'error': str(val_err)}), 400
    except Exception as err:
        return jsonify({'error': f"Internal Process Halt: {str(err)}"}), 500


if __name__ == "__main__":
    print_banner()

    # If terminal flags are supplied, process via command line interface
    if len(sys.argv) > 1 and sys.argv[1] not in ['run', 'serve']:
        parser = argparse.ArgumentParser(
            description="Fetch complete book data from Usul.ai API and generate a Word Document (.docx) directly."
        )
        parser.add_argument("url", type=str, help="The Usul.ai book URL")
        parser.add_argument("-o", "--output", type=str, default="book.docx", help="Output path")
        args = parser.parse_args()

        try:
            api_url = clean_and_format_url(args.url)
            complete_book_data = fetch_entire_book(api_url, chunk_size=50)
            if complete_book_data:
                doc = build_docx_from_data(complete_book_data)
                if doc:
                    doc.save(args.output)
                    print(f"\r{STATUS_OK} Document compiled successfully: {CLR_BOLD}{CLR_GREEN}{args.output}{CLR_RESET}\n")
                else:
                    print(f"{STATUS_WARN} Missing valid book text tags.")
            else:
                print(f"{STATUS_ERROR} {CLR_FAIL}Operation failed: Execution pipeline empty.{CLR_RESET}\n")
        except Exception as e:
            print(f"{STATUS_ERROR} Error: {e}")
            sys.exit(1)
    else:
        # Fallback default: Boot interactive Flask framework engine instance
        print(f"{STATUS_INFO} Booting Interactive Web UI interface server...")
        print(f"{STATUS_INFO} Open browser and head to: {CLR_BOLD}{CLR_GREEN}http://127.0.0.1:5000{CLR_RESET}")
        app.run(debug=True, port=5000)
